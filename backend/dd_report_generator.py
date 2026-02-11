# dd_report_generator.py
from docx import Document

def generate_dd_report(dd_summary: dict, output_file: str):
    """
    Generates a basic Due Diligence report (DOCX).
    SAFE placeholder implementation.
    Does NOT change any schemas or logic.
    """

    doc = Document()
    doc.add_heading("Due Diligence Report", level=1)

    if not dd_summary:
        doc.add_paragraph("No Due Diligence data available.")
        doc.save(output_file)
        return

    # Document summaries
    documents = dd_summary.get("documents", {})
    for doc_name, data in documents.items():
        doc.add_heading(doc_name, level=2)
        doc.add_paragraph(f"Document Type: {data.get('doc_type')}")
        doc.add_paragraph(f"Overall Risk: {data.get('overall_risk')}")

        counts = data.get("risk_counts", {})
        doc.add_paragraph(
            f"High: {counts.get('High', 0)}, "
            f"Medium: {counts.get('Medium', 0)}, "
            f"Low: {counts.get('Low', 0)}, "
            f"Total: {data.get('total_risks', 0)}"
        )

    doc.save(output_file)
